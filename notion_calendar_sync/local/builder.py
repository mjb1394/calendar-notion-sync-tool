"""
README
======

This script generates a printable monthly calendar as PDF and Word (DOCX)
files using events from a JSON file.

Installation:
    pip install reportlab python-docx python-dateutil babel

Usage:
    python calendar_builder.py \
        --events events.json \
        --year 2025 \
        --month 9 \
        --view monthly \
        --out calendar_sep_2025 \
        --format pdf docx \
        --page A4 \
        --locale en_US \
        --timezone UTC

Sample JSON:
[
  {"event":"Labor Day","eventtype":"holiday","location":"","start":"2025-09-01"},
  {"event":"Board Meeting","eventtype":"meeting","location":"HQ Rm 3","start":"2025-09-12T09:00","endtime":"2025-09-12T10:30"}
]
"""

from __future__ import annotations

import argparse
import calendar as stdlib_calendar
import logging
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

from babel.dates import format_date, format_datetime
from dateutil import parser as dateparser
from zoneinfo import ZoneInfo

# PDF and DOCX generation libraries
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter, landscape
from reportlab.pdfgen import canvas

from notion_calendar_sync.utils import JsonManager

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")


@dataclass
class Event:
    """Representation of a calendar event."""

    title: str
    eventtype: str
    location: str
    start: datetime
    end: datetime
    all_day: bool = False

    def display_title(self, cont: bool = False) -> str:
        suffix = " (cont.)" if cont else ""
        if self.location:
            return f"{self.title} ({self.location}){suffix}"
        return f"{self.title}{suffix}"


@dataclass
class Task:
    title: str
    due_date: date
    priority: str
    notes: str = ""


EVENT_COLORS = {
    "meeting": colors.blue,
    "holiday": colors.green,
    "deadline": colors.red,
    "default": colors.gray,
}


def parse_items(path: Path, timezone: ZoneInfo) -> Tuple[List[Event], List[Task]]:
    """Parse events and tasks from a JSON file."""

    events: List[Event] = []
    tasks: List[Task] = []
    json_manager = JsonManager()
    items = json_manager.read_json_file(path)

    for lineno, data in enumerate(items, 1):
        if data.get("type") == "task" or "task" in data:
            try:
                due = dateparser.isoparse(data.get("due_date")).date()
            except Exception:
                logging.warning("Invalid task item %s: missing due_date", lineno)
                continue
            task = Task(
                title=data.get("task", "").strip(),
                due_date=due,
                priority=data.get("priority", "medium").strip(),
                notes=data.get("notes", "").strip(),
            )
            tasks.append(task)
            continue

        # Event parsing
        title = data.get("event", "").strip()
        if not title:
            logging.warning("Skipping item %s: missing event title", lineno)
            continue
        eventtype = data.get("eventtype", "default").strip()
        location = data.get("location", "").strip()
        room = data.get("room")
        if room:
            location = f"{location} {room}".strip()

        if "date" in data:
            start_str = f"{data['date']}T{data.get('start', '00:00')}"
        else:
            start_str = data.get("start", "")
        if not start_str:
            logging.warning("Skipping item %s: missing start", lineno)
            continue
        start_dt = dateparser.isoparse(start_str)
        is_all_day = (
            "T" not in start_str
            or start_str.endswith("T00:00")
            and data.get("end") is None
            and data.get("endtime") is None
        )
        if start_dt.tzinfo is None:
            start_dt = start_dt.replace(tzinfo=timezone)
        start_dt = start_dt.astimezone(timezone)

        end_str = data.get("endtime")
        if not end_str and "end" in data:
            end_str = f"{data['date']}T{data['end']}"
        if end_str:
            end_dt = dateparser.isoparse(end_str)
            if end_dt.tzinfo is None:
                end_dt = end_dt.replace(tzinfo=timezone)
            end_dt = end_dt.astimezone(timezone)
        else:
            if is_all_day:
                end_dt = start_dt + timedelta(days=1)
            else:
                end_dt = start_dt + timedelta(hours=1)
            logging.info("Inferred end time for '%s'", title)

        events.append(
            Event(title, eventtype, location, start_dt, end_dt, is_all_day)
        )

    return events, tasks


def parse_events(path: Path, timezone: ZoneInfo) -> List[Event]:
    """Backward-compatible event parser."""

    events, _ = parse_items(path, timezone)
    return events


def month_matrix(year: int, month: int, week_start: str) -> List[List[date]]:
    """Return a matrix (weeks x days) for the month."""

    week_start_map = {
        "Sunday": stdlib_calendar.SUNDAY,
        "Monday": stdlib_calendar.MONDAY,
    }
    start = week_start_map.get(week_start, stdlib_calendar.SUNDAY)
    cal = stdlib_calendar.Calendar(firstweekday=start)
    matrix = cal.monthdatescalendar(year, month)
    return matrix


def week_matrix(week_date: date, week_start: str) -> List[date]:
    """Return list of 7 dates for week containing week_date."""
    week_start_map = {
        "Sunday": stdlib_calendar.SUNDAY,
        "Monday": stdlib_calendar.MONDAY,
    }
    start_idx = week_start_map.get(week_start, stdlib_calendar.SUNDAY)
    day = week_date
    while day.weekday() != start_idx:
        day -= timedelta(days=1)
    return [day + timedelta(days=i) for i in range(7)]


def group_events_by_day(events: Iterable[Event]) -> Dict[date, List[Event]]:
    """Group events by their date, expanding multi-day events."""

    grouped: Dict[date, List[Event]] = {}
    for ev in events:
        day = ev.start.date()
        end_day = ev.end.date()
        while day <= end_day:
            grouped.setdefault(day, []).append(ev)
            day += timedelta(days=1)
    for day_events in grouped.values():
        day_events.sort(key=lambda e: e.start)
    return grouped


def filter_tasks(
    tasks: List[Task], task_filter: str, start: date, end: date
) -> List[Task]:
    """Filter tasks according to filter string and date range."""
    filtered = tasks
    if task_filter == "due":
        filtered = [t for t in tasks if start <= t.due_date <= end]
    elif task_filter == "overdue":
        today = date.today()
        filtered = [t for t in tasks if t.due_date < today]
    elif task_filter.startswith("priority="):
        level = task_filter.split("=", 1)[1]
        filtered = [t for t in tasks if t.priority == level]
    return sorted(
        filtered,
        key=lambda t: (
            t.due_date,
            {"high": 0, "medium": 1, "low": 2}.get(t.priority, 1),
        ),
    )


def layout_day_cell(
    day: date, events: List[Event], tz: ZoneInfo, loc: str
) -> List[str]:
    """Prepare strings for a day cell."""
    lines: List[str] = []
    lines.append(str(day.day))
    for ev in events[:5]:
        if ev.all_day or (
            ev.start.hour == 0
            and ev.start.minute == 0
            and ev.end - ev.start >= timedelta(days=1)
        ):
            line = ev.display_title(day > ev.start.date())
        else:
            start = format_datetime(ev.start.astimezone(tz), "HH:mm", locale=loc)
            end = format_datetime(ev.end.astimezone(tz), "HH:mm", locale=loc)
            line = f"{start}-{end} {ev.display_title(day > ev.start.date())}"
        if len(line) > 40:
            line = line[:37] + "..."
        lines.append(line)
    if len(events) > 5:
        lines.append(f"+{len(events) - 5} more")
    return lines


def render_tasks_pdf(c: canvas.Canvas, tasks: List[Task], page_size) -> None:
    """Render tasks on the current PDF page."""
    width, height = page_size
    margin = 0.5 * 72
    y = height - margin
    c.setFont("Helvetica-Bold", 16)
    c.drawString(margin, y, "Tasks")
    y -= 24
    c.setFont("Helvetica-Bold", 12)
    headers = ["Task", "Due", "Priority", "Notes"]
    col_width = (width - 2 * margin) / 4
    for i, h in enumerate(headers):
        c.drawString(margin + i * col_width, y, h)
    y -= 14
    c.setFont("Helvetica", 11)
    for t in tasks:
        if y < margin:
            c.showPage()
            y = height - margin
        c.drawString(margin, y, t.title)
        c.drawString(margin + col_width, y, t.due_date.isoformat())
        c.drawString(margin + 2 * col_width, y, t.priority)
        c.drawString(margin + 3 * col_width, y, t.notes)
        y -= 14


def render_tasks_docx(doc: Document, tasks: List[Task]) -> None:
    doc.add_heading("Tasks", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Task"
    hdr[1].text = "Due"
    hdr[2].text = "Priority"
    hdr[3].text = "Notes"
    for t in tasks:
        row = table.add_row().cells
        row[0].text = t.title
        row[1].text = t.due_date.isoformat()
        row[2].text = t.priority
        row[3].text = t.notes


def render_weekly_pdf(
    output: Path,
    days: List[date],
    events: Dict[date, List[Event]],
    page: str,
    tz: ZoneInfo,
    loc: str,
    tasks: List[Task] | None = None,
) -> None:
    """Render weekly calendar to PDF with optional tasks page."""
    page_sizes = {"A4": A4, "Letter": letter}
    page_size = landscape(page_sizes.get(page, letter))
    c = canvas.Canvas(str(output), pagesize=page_size)
    width, height = page_size
    margin = 0.5 * 72

    cell_width = (width - 2 * margin) / 7
    cell_height = height - 2 * margin
    c.setFont("Helvetica-Bold", 14)
    week_label = f"{days[0].isoformat()} - {days[-1].isoformat()}"
    c.drawCentredString(width / 2, height - margin / 2, week_label)

    for idx, day in enumerate(days):
        x = margin + idx * cell_width
        y = height - margin - cell_height
        c.rect(x, y, cell_width, cell_height)
        day_events = events.get(day, [])
        lines = layout_day_cell(day, day_events, tz, loc)
        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(x + cell_width - 2, y + cell_height - 12, lines[0])
        c.setFont("Helvetica", 9)
        for i, text in enumerate(lines[1:], 1):
            color = EVENT_COLORS.get(
                day_events[i - 1].eventtype if i - 1 < len(day_events) else "default",
                colors.gray,
            )
            c.setFillColor(color)
            c.drawString(x + 2, y + cell_height - 12 - i * 10, text)
            c.setFillColor(colors.black)

    if tasks:
        c.showPage()
        render_tasks_pdf(c, tasks, page_size)
    c.save()
    logging.info("Wrote %s", output)


def render_weekly_docx(
    output: Path,
    days: List[date],
    events: Dict[date, List[Event]],
    page: str,
    tz: ZoneInfo,
    loc: str,
    tasks: List[Task] | None = None,
) -> None:
    """Render weekly calendar to DOCX with optional tasks page."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    if page == "A4":
        section.page_height, section.page_width = (
            section.page_width,
            section.page_height,
        )
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)

    table = doc.add_table(rows=1, cols=7)
    widths = [Inches((section.page_width.inches - 1) / 7) for _ in range(7)]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    row = table.rows[0]
    for day_idx, day in enumerate(days):
        cell = row.cells[day_idx]
        lines = layout_day_cell(day, events.get(day, []), tz, loc)
        p = cell.paragraphs[0]
        run = p.add_run(lines[0])
        run.bold = True
        p.alignment = 2
        for text in lines[1:]:
            p = cell.add_paragraph()
            p.alignment = 0
            run = p.add_run(text)
            run.font.size = Pt(9)

    if tasks:
        doc.add_page_break()
        render_tasks_docx(doc, tasks)
    doc.save(str(output))
    logging.info("Wrote %s", output)


def render_monthly_pdf(
    output: Path,
    matrix: List[List[date]],
    events: Dict[date, List[Event]],
    year: int,
    month: int,
    page: str,
    tz: ZoneInfo,
    loc: str,
    tasks: List[Task] | None = None,
) -> None:
    """Render monthly calendar to a PDF file. Optionally append tasks page."""

    page_sizes = {"A4": A4, "Letter": letter}
    page_size = landscape(page_sizes.get(page, letter))
    c = canvas.Canvas(str(output), pagesize=page_size)
    width, height = page_size
    margin = 0.5 * 72  # 0.5 inch

    cell_width = (width - 2 * margin) / 7
    cell_height = (height - 2 * margin) / len(matrix)

    month_name = format_date(date(year, month, 1), "LLLL yyyy", locale=loc)
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, height - margin / 2, month_name)

    for week_idx, week in enumerate(matrix):
        for day_idx, day in enumerate(week):
            x = margin + day_idx * cell_width
            y = height - margin - (week_idx + 1) * cell_height
            c.rect(x, y, cell_width, cell_height)
            day_events = events.get(day, [])
            lines = layout_day_cell(day, day_events, tz, loc)
            c.setFont("Helvetica-Bold", 10)
            c.drawRightString(x + cell_width - 2, y + cell_height - 12, lines[0])
            c.setFont("Helvetica", 9)
            for i, text in enumerate(lines[1:], 1):
                color = EVENT_COLORS.get(
                    (
                        day_events[i - 1].eventtype
                        if i - 1 < len(day_events)
                        else "default"
                    ),
                    colors.gray,
                )
                c.setFillColor(color)
                c.drawString(x + 2, y + cell_height - 12 - i * 10, text)
                c.setFillColor(colors.black)

    # Legend
    legend_y = margin / 2
    c.setFont("Helvetica", 9)
    x = margin
    for key, color in EVENT_COLORS.items():
        c.setFillColor(color)
        c.rect(x, legend_y, 10, 10, fill=1)
        c.setFillColor(colors.black)
        c.drawString(x + 14, legend_y, key)
        x += 60

    if tasks:
        c.showPage()
        render_tasks_pdf(c, tasks, page_size)
    c.save()
    logging.info("Wrote %s", output)


def render_monthly_docx(
    output: Path,
    matrix: List[List[date]],
    events: Dict[date, List[Event]],
    year: int,
    month: int,
    page: str,
    tz: ZoneInfo,
    loc: str,
    tasks: List[Task] | None = None,
) -> None:
    """Render monthly calendar to a DOCX file. Optionally append tasks page."""

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    if page == "A4":
        section.page_height, section.page_width = (
            section.page_width,
            section.page_height,
        )
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)

    table = doc.add_table(rows=len(matrix), cols=7)
    widths = [Inches((section.page_width.inches - 1) / 7) for _ in range(7)]
    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    for week_idx, week in enumerate(matrix):
        for day_idx, day in enumerate(week):
            cell = table.rows[week_idx].cells[day_idx]
            lines = layout_day_cell(day, events.get(day, []), tz, loc)
            p = cell.paragraphs[0]
            run = p.add_run(lines[0])
            run.bold = True
            p.alignment = 2  # right
            for text in lines[1:]:
                p = cell.add_paragraph()
                p.alignment = 0
                run = p.add_run(text)
                run.font.size = Pt(9)

    # Legend
    doc.add_paragraph("Legend:")
    for key, color in EVENT_COLORS.items():
        para = doc.add_paragraph()
        run = para.add_run("\u25a0 ")
        run.font.color.rgb = RGBColor(
            int(color.red * 255), int(color.green * 255), int(color.blue * 255)
        )
        para.add_run(key)

    if tasks:
        doc.add_page_break()
        render_tasks_docx(doc, tasks)
    doc.save(str(output))
    logging.info("Wrote %s", output)


def build_calendar(
    events_path: Path,
    view: str,
    year: int | None = None,
    month: int | None = None,
    week: date | None = None,
    week_start: str = "Sunday",
    include_tasks: bool = False,
    tasks_filter: str = "all",
    out_prefix: str = "cal_",
    formats: List[str] | None = None,
    page: str = "Letter",
    locale: str = "en_US",
    timezone: str = "UTC",
) -> List[Path]:
    """Build calendar files and return list of output paths."""

    formats = formats or ["pdf"]
    tz = ZoneInfo(timezone)
    events, tasks = parse_items(events_path, tz)
    outputs: List[Path] = []

    if view == "monthly":
        if year is None or month is None:
            raise ValueError("year and month required for monthly view")
        matrix = month_matrix(year, month, week_start)
        grouped = group_events_by_day(events)
        start = date(year, month, 1)
        end = (start + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        task_list = (
            filter_tasks(tasks, tasks_filter, start, end) if include_tasks else None
        )
        base = f"{out_prefix}{year:04d}-{month:02d}"
        if "pdf" in formats:
            out = Path(f"{base}.pdf")
            render_monthly_pdf(
                out, matrix, grouped, year, month, page, tz, locale, task_list
            )
            outputs.append(out)
        if "docx" in formats:
            out = Path(f"{base}.docx")
            render_monthly_docx(
                out, matrix, grouped, year, month, page, tz, locale, task_list
            )
            outputs.append(out)
    else:  # weekly
        if week is None:
            raise ValueError("week date required for weekly view")
        days = week_matrix(week, week_start)
        grouped = group_events_by_day(events)
        start = days[0]
        end = days[-1]
        task_list = (
            filter_tasks(tasks, tasks_filter, start, end) if include_tasks else None
        )
        base = f"{out_prefix}{start.isoformat()}_to_{end.isoformat()}"
        if "pdf" in formats:
            out = Path(f"{base}.pdf")
            render_weekly_pdf(out, days, grouped, page, tz, locale, task_list)
            outputs.append(out)
        if "docx" in formats:
            out = Path(f"{base}.docx")
            render_weekly_docx(out, days, grouped, page, tz, locale, task_list)
            outputs.append(out)

    return outputs


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate calendar PDF/DOCX from events"
    )
    parser.add_argument("--events", required=True, help="Path to events JSON file")
    parser.add_argument(
        "--view", default="monthly", choices=["monthly", "weekly"], help="Calendar view"
    )
    parser.add_argument("--year", type=int, help="Year for monthly view")
    parser.add_argument("--month", type=int, help="Month (1-12) for monthly view")
    parser.add_argument("--week", help="ISO date within desired week for weekly view")
    parser.add_argument(
        "--week-start",
        default="Sunday",
        choices=["Sunday", "Monday"],
        help="First day of the week",
    )
    parser.add_argument(
        "--include-tasks", action="store_true", help="Include tasks page"
    )
    parser.add_argument("--tasks-filter", default="all", help="Filter for tasks")
    parser.add_argument(
        "--out-prefix", default="cal_", help="Prefix for output filenames"
    )
    parser.add_argument(
        "--format",
        nargs="+",
        default=["pdf"],
        choices=["pdf", "docx"],
        help="Output formats",
    )
    parser.add_argument(
        "--page", default="Letter", choices=["A4", "Letter"], help="Page size"
    )
    parser.add_argument(
        "--locale", default="en_US", help="Locale for month and weekdays"
    )
    parser.add_argument("--timezone", default="UTC", help="Timezone for events")

    args = parser.parse_args()

    week_date = dateparser.isoparse(args.week).date() if args.week else None
    outputs = build_calendar(
        events_path=Path(args.events),
        view=args.view,
        year=args.year,
        month=args.month,
        week=week_date,
        week_start=args.week_start,
        include_tasks=args.include_tasks,
        tasks_filter=args.tasks_filter,
        out_prefix=args.out_prefix,
        formats=args.format,
        page=args.page,
        locale=args.locale,
        timezone=args.timezone,
    )
    for out in outputs:
        logging.info("Wrote %s", out)


if __name__ == "__main__":
    main()
